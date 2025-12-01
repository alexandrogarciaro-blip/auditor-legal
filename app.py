import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import time
import tempfile
import os
import re
from datetime import datetime

# --- 1. CONFIGURACIÓN VISUAL ---
st.set_page_config(page_title="LegalAudit AI", page_icon="⚖️", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    /* BARRA LATERAL OSCURA */
    section[data-testid="stSidebar"] {background-color: #101820;}
    section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] label, section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] p {color: #ffffff !important;}
    
    /* FONDO PRINCIPAL */
    .main {background-color: #f4f6f9;}
    h1 {color: #2c3e50; font-family: 'Helvetica', sans-serif;}
    
    /* BOTONES DORADOS */
    .stButton>button {width: 100%; border-radius: 8px; height: 3em; background-color: #c5a059; color: white; font-weight: bold; border: none;}
    .stButton>button:hover {background-color: #b08d4b; color: white;}
    
    /* VISIBILIDAD DE ARCHIVOS EN BARRA LATERAL */
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] div,
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] small,
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] span {color: #ffffff !important;}
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] svg {fill: #ffffff !important;}
    [data-testid="stSidebar"] button[kind="secondary"] {background-color: #ffffff !important; color: #000000 !important; border: none;}
    
    /* CAJA DE ÉXITO */
    .success-box {padding: 1rem; background-color: #d4edda; border-left: 6px solid #28a745; color: #155724; margin-bottom: 1rem;}
    </style>
    """, unsafe_allow_html=True)

# --- 2. CONEXIÓN ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except:
    st.error("⚠️ Error: No API Key found.")
    st.stop()

# --- 3. FUNCIONES DE LIMPIEZA INTELIGENTE (RESCATE DE TABLAS) ---

def clean_technical_output(text):
    """
    Estrategia de limpieza V5.6:
    1. Quita los envoltorios de código (```).
    2. Filtra línea a línea borrando lo que parece Python.
    3. PROTEGE explícitamente las líneas que parecen tablas Markdown (|...|).
    """
    # 1. Quitar envoltorios pero dejar el contenido
    text = text.replace("```markdown", "").replace("```python", "").replace("```", "")
    
    # 2. Corregir título si hace falta
    text = text.replace("# INFORME DE DUE DILIGENCE", "# INFORME DE SITUACIÓN")
    
    lines = text.split('\n')
    clean_lines = []
    
    for line in lines:
        l = line.strip()
        
        # --- REGLA DE PROTECCIÓN DE TABLAS ---
        # Si la línea empieza y acaba por '|', es una tabla. LA GUARDAMOS SIEMPRE.
        if l.startswith("|") and l.endswith("|"):
            clean_lines.append(line)
            continue
        
        # --- FILTRO DE BASURA TÉCNICA ---
        is_garbage = False
        
        # Detectar sintaxis Python común
        if l.startswith("print(") or l.startswith("def ") or l.startswith("import "): is_garbage = True
        if " = {" in l or " = [" in l: is_garbage = True # Asignación de variables
        if "append(" in l or "return " in l: is_garbage = True
        if l == "python": is_garbage = True
        
        # Si no es basura, lo guardamos (es texto narrativo)
        if not is_garbage:
            clean_lines.append(line)
            
    return '\n'.join(clean_lines).strip()

def add_markdown_to_doc(doc, text):
    lines = text.split('\n')
    table_buffer = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        if stripped.startswith('|') and stripped.endswith('|'):
            if '---' in stripped: continue
            row_data = [c.strip() for c in stripped.split('|') if c.strip()]
            table_buffer.append(row_data)
            in_table = True
        else:
            if in_table and table_buffer:
                if len(table_buffer) > 0:
                    rows = len(table_buffer)
                    cols = len(table_buffer[0])
                    t = doc.add_table(rows=rows, cols=cols)
                    t.style = 'Table Grid'
                    t.autofit = True
                    for r, row_data in enumerate(table_buffer):
                        for c, cell_text in enumerate(row_data):
                            if c < cols:
                                cell = t.cell(r, c)
                                p = cell.paragraphs[0]
                                p.text = cell_text
                                
                                is_header = (r == 0)
                                is_total = (c == 0 and "TOTAL" in cell_text.upper())
                                
                                if is_header or is_total: 
                                    for run in p.runs: run.bold = True
                                if "TOTAL" in row_data[0].upper():
                                     for run in p.runs: run.bold = True
                table_buffer = []
                in_table = False

            if stripped.startswith('## '):
                doc.add_heading(stripped.replace('#', '').strip(), level=1)
            elif stripped.startswith('### '):
                doc.add_heading(stripped.replace('#', '').strip(), level=2)
            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
    return doc

def create_professional_report(content_text):
    doc = Document()
    for _ in range(5): doc.add_paragraph()
    title = doc.add_heading('INFORME DE SITUACIÓN SOCIETARIA', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    add_markdown_to_doc(doc, content_text)
    return doc

# --- 4. INTERFAZ ---
with st.sidebar:
    try:
        st.image("logo.png", width=280)
    except:
        st.image("https://cdn-icons-png.flaticon.com/512/1998/1998342.png", width=100)
    st.markdown("### Panel de Control")
    uploaded_files = st.file_uploader("1. Sube Escrituras (PDF)", type=['pdf'], accept_multiple_files=True)
    st.markdown("---")
    analyze_btn = st.button("2. EJECUTAR ANÁLISIS ✨", type="primary")
    st.markdown("---")
    st.info("💡 **Consejo:** Sube todos los documentos de una misma empresa juntos.")

# --- 5. INTERFAZ PRINCIPAL ---
st.title("⚖️ Auditoría Legal Inteligente")

if not uploaded_files:
    st.markdown("""
    <div style="padding: 20px; background-color: #e8f4f8; border-radius: 10px; border: 1px solid #d1e7dd;">
        <h4 style="color: #0c5460;">👋 Bienvenido al Sistema de Auditoría</h4>
        <p style="color: #0c5460;">Herramienta avanzada para analizar escrituras notariales.</p>
        <p><b>Instrucciones:</b></p>
        <ol style="color: #0c5460;">
            <li>Sube los PDFs en el menú de la izquierda.</li>
            <li>Haz clic en <b>EJECUTAR ANÁLISIS</b>.</li>
            <li>La IA ordenará los hechos y calculará el reparto de capital.</li>
            <li>Podrás descargar el resultado en Word.</li>
        </ol>
    </div>
    """, unsafe_allow_html=True)

if analyze_btn and uploaded_files:
    tab1, tab2 = st.tabs(["📄 Informe", "📥 Word"])
    
    with tab1:
        progress = st.progress(0, text="Procesando...")
        try:
            gemini_files = []
            for i, f in enumerate(uploaded_files):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(f.getvalue())
                    tmp_path = tmp.name
                g_file = genai.upload_file(path=tmp_path, display_name=f.name)
                gemini_files.append(g_file)
                os.remove(tmp_path)

            progress.progress(0.6, text="Analizando...")
            time.sleep(1)
            
            # --- PROMPT V5.6 (MÁS NARRATIVO PERO ESTRICTO CON LA TABLA) ---
            SYSTEM_PROMPT = """
            ROL: Abogado Mercantilista y Auditor.
            OBJETIVO: Redactar un Informe de Situación Societaria.
            
            INSTRUCCIONES TÉCNICAS:
            1. Usa 'code_execution' para calcular el Cap Table.
            2. LA TABLA ES OBLIGATORIA. Si la calculas en Python, imprímela también en formato Markdown (|...|) para que se vea en el informe.
            
            FORMATO DE TABLA OBLIGATORIO:
            | Socios | Participaciones | Capital Nominal | Porcentaje % |
            |---|---|---|---|
            | [Datos...] | [Datos...] | [Datos...] | [Datos...] |
            | **TOTAL** | **[Suma]** | **[Suma]** | **100%** |
            
            ESTRUCTURA DEL INFORME:
            1. Resumen Ejecutivo.
            2. Cronología Detallada.
            3. Tabla de Titularidad Actual (OBLIGATORIA).
            4. Incidencias.
            """

            # Temperatura baja para consistencia
            generation_config = {"temperature": 0.1}

            model = genai.GenerativeModel(
                model_name="gemini-2.5-flash",
                system_instruction=SYSTEM_PROMPT,
                generation_config=generation_config,
                tools='code_execution'
            )
            response = model.generate_content(["Genera el informe.", *gemini_files])
            
            # APLICAMOS LA NUEVA LIMPIEZA INTELIGENTE
            final_text = clean_technical_output(response.text)
            
            progress.empty()
            st.markdown('<div class="success-box">✅ Análisis completado.</div>', unsafe_allow_html=True)
            st.markdown(final_text)
            st.session_state['report_text'] = final_text

        except Exception as e:
            st.error(f"Error: {e}")

    with tab2:
        if 'report_text' in st.session_state:
            st.write("Descarga el documento final.")
            doc = create_professional_report(st.session_state['report_text'])
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("📥 Descargar Word", data=bio.getvalue(), file_name="Auditoria.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
